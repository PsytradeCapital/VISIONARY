"""
Advanced document parsing with cloud-based NLP processing
Task 3.1: Enhanced document parsing with cloud processing
"""

import asyncio
import io
from typing import Tuple, Dict, Any, Optional
import PyPDF2
import docx
from PIL import Image
import pytesseract
import fitz  # PyMuPDF for better PDF handling
from pydantic import BaseModel

class DocumentParser:
    """Advanced document parser with cloud-based NLP processing"""
    
    def __init__(self):
        # OCR settings for image text extraction
        self.ocr_config = '--oem 3 --psm 6'
        
        # Supported document formats
        self.supported_formats = {
            'pdf': ['.pdf'],
            'word': ['.docx', '.doc'],
            'text': ['.txt', '.rtf'],
            'image': ['.jpg', '.jpeg', '.png', '.bmp', '.tiff']
        }
    
    async def parse_pdf(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """
        Parse PDF with advanced text extraction and metadata
        Uses PyMuPDF for better text extraction than PyPDF2
        """
        try:
            # Use PyMuPDF for better text extraction
            doc = fitz.open(stream=file_content, filetype="pdf")
            
            extracted_text = ""
            metadata = {
                "page_count": len(doc),
                "has_images": False,
                "has_tables": False,
                "text_blocks": [],
                "images_extracted": 0,
                "processing_method": "pymupdf"
            }
            
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                
                # Extract text
                page_text = page.get_text()
                extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
                
                # Extract text blocks with positioning
                text_blocks = page.get_text("dict")
                metadata["text_blocks"].append({
                    "page": page_num + 1,
                    "blocks": len(text_blocks.get("blocks", []))
                })
                
                # Check for images
                image_list = page.get_images()
                if image_list:
                    metadata["has_images"] = True
                    metadata["images_extracted"] += len(image_list)
                
                # Check for tables (basic detection)
                if "table" in page_text.lower() or "|" in page_text:
                    metadata["has_tables"] = True
            
            doc.close()
            
            # Fallback to PyPDF2 if PyMuPDF fails or returns empty text
            if not extracted_text.strip():
                extracted_text, fallback_metadata = await self._parse_pdf_pypdf2(file_content)
                metadata.update(fallback_metadata)
                metadata["processing_method"] = "pypdf2_fallback"
            
            return extracted_text.strip(), metadata
            
        except Exception as e:
            # Fallback to PyPDF2
            try:
                return await self._parse_pdf_pypdf2(file_content)
            except Exception as fallback_error:
                return f"Error parsing PDF: {str(e)}, Fallback error: {str(fallback_error)}", {
                    "error": True,
                    "error_message": str(e),
                    "fallback_error": str(fallback_error)
                }
    
    async def _parse_pdf_pypdf2(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Fallback PDF parsing using PyPDF2"""
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        
        extracted_text = ""
        metadata = {
            "page_count": len(pdf_reader.pages),
            "processing_method": "pypdf2",
            "encrypted": pdf_reader.is_encrypted
        }
        
        # Extract metadata
        if pdf_reader.metadata:
            metadata.update({
                "title": pdf_reader.metadata.get("/Title", ""),
                "author": pdf_reader.metadata.get("/Author", ""),
                "subject": pdf_reader.metadata.get("/Subject", ""),
                "creator": pdf_reader.metadata.get("/Creator", "")
            })
        
        # Extract text from all pages
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                extracted_text += f"\n--- Page {page_num + 1} ---\n{page_text}\n"
            except Exception as e:
                extracted_text += f"\n--- Page {page_num + 1} (Error: {str(e)}) ---\n"
        
        return extracted_text.strip(), metadata
    
    async def parse_docx(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Parse DOCX files with metadata extraction"""
        try:
            doc = docx.Document(io.BytesIO(file_content))
            
            extracted_text = ""
            metadata = {
                "paragraph_count": 0,
                "table_count": 0,
                "image_count": 0,
                "style_info": {},
                "processing_method": "python-docx"
            }
            
            # Extract core properties
            if doc.core_properties:
                metadata.update({
                    "title": doc.core_properties.title or "",
                    "author": doc.core_properties.author or "",
                    "subject": doc.core_properties.subject or "",
                    "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                    "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                })
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    extracted_text += paragraph.text + "\n"
                    metadata["paragraph_count"] += 1
                    
                    # Track styles
                    style_name = paragraph.style.name
                    if style_name in metadata["style_info"]:
                        metadata["style_info"][style_name] += 1
                    else:
                        metadata["style_info"][style_name] = 1
            
            # Extract tables
            for table in doc.tables:
                metadata["table_count"] += 1
                extracted_text += "\n--- Table ---\n"
                
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    extracted_text += " | ".join(row_text) + "\n"
                
                extracted_text += "--- End Table ---\n"
            
            # Count images (inline shapes)
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    metadata["image_count"] += 1
            
            return extracted_text.strip(), metadata
            
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}", {
                "error": True,
                "error_message": str(e)
            }
    
    async def extract_text_from_image(self, file_content: bytes) -> Tuple[str, Dict[str, Any]]:
        """Extract text from images using OCR (Tesseract)"""
        try:
            # Open image
            image = Image.open(io.BytesIO(file_content))
            
            metadata = {
                "image_format": image.format,
                "image_size": image.size,
                "image_mode": image.mode,
                "processing_method": "tesseract_ocr"
            }
            
            # Convert to RGB if necessary
            if image.mode != 'RGB':
                image = image.convert('RGB')
            
            # Perform OCR
            extracted_text = pytesseract.image_to_string(image, config=self.ocr_config)
            
            # Get OCR confidence data
            try:
                ocr_data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
                confidences = [int(conf) for conf in ocr_data['conf'] if int(conf) > 0]
                
                if confidences:
                    metadata.update({
                        "ocr_confidence_avg": sum(confidences) / len(confidences),
                        "ocr_confidence_min": min(confidences),
                        "ocr_confidence_max": max(confidences),
                        "words_detected": len([word for word in ocr_data['text'] if word.strip()])
                    })
            except Exception as ocr_error:
                metadata["ocr_confidence_error"] = str(ocr_error)
            
            return extracted_text.strip(), metadata
            
        except Exception as e:
            return f"Error extracting text from image: {str(e)}", {
                "error": True,
                "error_message": str(e)
            }
    
    async def parse_text_file(self, file_content: bytes, encoding: str = 'utf-8') -> Tuple[str, Dict[str, Any]]:
        """Parse plain text files with encoding detection"""
        try:
            # Try specified encoding first
            try:
                text = file_content.decode(encoding)
            except UnicodeDecodeError:
                # Try common encodings
                encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
                text = None
                used_encoding = None
                
                for enc in encodings:
                    try:
                        text = file_content.decode(enc)
                        used_encoding = enc
                        break
                    except UnicodeDecodeError:
                        continue
                
                if text is None:
                    # Last resort - decode with errors ignored
                    text = file_content.decode('utf-8', errors='ignore')
                    used_encoding = 'utf-8 (with errors ignored)'
            else:
                used_encoding = encoding
            
            # Analyze text
            lines = text.split('\n')
            words = text.split()
            
            metadata = {
                "encoding_used": used_encoding,
                "line_count": len(lines),
                "word_count": len(words),
                "character_count": len(text),
                "empty_lines": len([line for line in lines if not line.strip()]),
                "processing_method": "text_decode"
            }
            
            return text, metadata
            
        except Exception as e:
            return f"Error parsing text file: {str(e)}", {
                "error": True,
                "error_message": str(e)
            }
    
    async def extract_metadata_only(self, file_content: bytes, file_type: str) -> Dict[str, Any]:
        """Extract only metadata without full text parsing (for quick analysis)"""
        try:
            if file_type.lower() == 'pdf':
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                metadata = {
                    "page_count": len(pdf_reader.pages),
                    "encrypted": pdf_reader.is_encrypted,
                    "file_type": "pdf"
                }
                
                if pdf_reader.metadata:
                    metadata.update({
                        "title": pdf_reader.metadata.get("/Title", ""),
                        "author": pdf_reader.metadata.get("/Author", ""),
                        "subject": pdf_reader.metadata.get("/Subject", ""),
                        "creator": pdf_reader.metadata.get("/Creator", "")
                    })
                
                return metadata
                
            elif file_type.lower() in ['docx', 'doc']:
                doc = docx.Document(io.BytesIO(file_content))
                metadata = {
                    "paragraph_count": len(doc.paragraphs),
                    "table_count": len(doc.tables),
                    "file_type": "docx"
                }
                
                if doc.core_properties:
                    metadata.update({
                        "title": doc.core_properties.title or "",
                        "author": doc.core_properties.author or "",
                        "subject": doc.core_properties.subject or "",
                        "created": str(doc.core_properties.created) if doc.core_properties.created else "",
                        "modified": str(doc.core_properties.modified) if doc.core_properties.modified else ""
                    })
                
                return metadata
                
            elif file_type.lower() in ['jpg', 'jpeg', 'png', 'bmp', 'tiff']:
                image = Image.open(io.BytesIO(file_content))
                metadata = {
                    "image_format": image.format,
                    "image_size": image.size,
                    "image_mode": image.mode,
                    "file_type": "image"
                }
                
                # Extract EXIF data if available
                if hasattr(image, '_getexif') and image._getexif():
                    exif_data = image._getexif()
                    metadata["exif_data"] = {k: str(v) for k, v in exif_data.items() if isinstance(v, (str, int, float))}
                
                return metadata
            
            else:
                return {
                    "file_type": file_type,
                    "size": len(file_content),
                    "metadata_extraction": "not_supported"
                }
                
        except Exception as e:
            return {
                "error": True,
                "error_message": str(e),
                "file_type": file_type
            }
    
    async def detect_document_language(self, text: str) -> Dict[str, Any]:
        """Detect document language (placeholder for actual implementation)"""
        # This would integrate with language detection libraries like:
        # - langdetect
        # - polyglot
        # - Google Cloud Translation API
        
        # Placeholder implementation
        common_english_words = ['the', 'and', 'or', 'but', 'in', 'on', 'at', 'to', 'for', 'of', 'with', 'by']
        text_lower = text.lower()
        english_word_count = sum(1 for word in common_english_words if word in text_lower)
        
        return {
            "detected_language": "en" if english_word_count > 3 else "unknown",
            "confidence": min(english_word_count / 10.0, 1.0),
            "method": "simple_keyword_detection"
        }
    
    async def extract_structured_data(self, text: str) -> Dict[str, Any]:
        """Extract structured data like dates, emails, phone numbers, etc."""
        import re
        
        # Regular expressions for common data types
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        phone_pattern = r'(\+\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}'
        date_pattern = r'\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b\d{4}[/-]\d{1,2}[/-]\d{1,2}\b'
        url_pattern = r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+'
        
        structured_data = {
            "emails": re.findall(email_pattern, text),
            "phone_numbers": re.findall(phone_pattern, text),
            "dates": re.findall(date_pattern, text),
            "urls": re.findall(url_pattern, text)
        }
        
        # Count occurrences
        structured_data["counts"] = {
            "emails": len(structured_data["emails"]),
            "phone_numbers": len(structured_data["phone_numbers"]),
            "dates": len(structured_data["dates"]),
            "urls": len(structured_data["urls"])
        }
        
        return structured_data